"""
Pure operation functions for PDF tools.
No HTTP, no Huey — just filesystem transforms.
Each function receives resolved Paths and returns Path(s).
"""
import io
import os
import shutil
import subprocess
import tempfile
import uuid
import zipfile
from pathlib import Path

import fitz  # PyMuPDF 1.27+
from PIL import Image

from common.executables import resolve_executable

# ── merge ─────────────────────────────────────────────────────────────────────

def merge_pdfs(input_paths: list[Path], output_path: Path) -> Path:
    doc = fitz.open()
    for p in input_paths:
        with fitz.open(p) as src:
            doc.insert_pdf(src)
    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    return output_path

# ── split ─────────────────────────────────────────────────────────────────────

def split_pdf(input_path: Path, output_dir: Path, ranges: list[tuple]) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    results = []
    with fitz.open(input_path) as doc:
        total = len(doc)
        effective = ranges if ranges else [(i + 1, i + 1) for i in range(total)]
        for idx, (start, end) in enumerate(effective):
            sub = fitz.open()
            sub.insert_pdf(doc, from_page=start - 1, to_page=min(end - 1, total - 1))
            name = f'part_{idx + 1:03d}.pdf' if ranges else f'page_{start:03d}.pdf'
            out = output_dir / name
            sub.save(out)
            sub.close()
            results.append(out)
    return results

# ── pdf → images ──────────────────────────────────────────────────────────────

def pdf_to_images(input_path: Path, output_dir: Path, dpi: int = 150) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    scale = dpi / 72
    mat = fitz.Matrix(scale, scale)
    results = []
    with fitz.open(input_path) as doc:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            out = output_dir / f'page_{i + 1:03d}.png'
            pix.save(str(out))
            results.append(out)
    return results

# ── images → pdf ──────────────────────────────────────────────────────────────

def images_to_pdf(input_paths: list[Path], output_path: Path) -> Path:
    doc = fitz.open()
    for img_path in input_paths:
        img = Image.open(img_path)
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        elif img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        w, h = img.size
        buf = io.BytesIO()
        img.save(buf, 'JPEG', quality=90)
        buf.seek(0)
        # Convert pixels to points (assume 96 DPI source)
        pw, ph = w * 72 / 96, h * 72 / 96
        page = doc.new_page(width=pw, height=ph)
        page.insert_image(page.rect, stream=buf.read())
    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    return output_path

# ── pdf → text ────────────────────────────────────────────────────────────────

def pdf_to_text(input_path: Path, output_path: Path) -> Path:
    with fitz.open(input_path) as doc:
        lines = []
        for i, page in enumerate(doc):
            lines.append(f'=== Página {i + 1} ===')
            lines.append(page.get_text())
    output_path.write_text('\n'.join(lines), encoding='utf-8')
    return output_path

# ── ocr image ─────────────────────────────────────────────────────────────────

def ocr_image(input_path: Path, output_path: Path, lang: str = 'por+eng') -> Path:
    tess = resolve_executable('tesseract')
    result = subprocess.run(
        [str(tess), str(input_path), 'stdout', '-l', lang],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or 'Tesseract retornou erro.')
    output_path.write_text(result.stdout, encoding='utf-8')
    return output_path

# ── ocr pdf ───────────────────────────────────────────────────────────────────

def ocr_pdf(input_path: Path, output_path: Path, lang: str = 'por+eng') -> Path:
    import ocrmypdf
    languages = lang.replace('+', ',').split(',')
    ocrmypdf.ocr(
        str(input_path),
        str(output_path),
        language=languages,
        skip_text=True,
        deskew=True,
        progress_bar=False,
    )
    return output_path

# ── compress ──────────────────────────────────────────────────────────────────

def compress_pdf(input_path: Path, output_path: Path, quality: str = 'screen') -> Path:
    """
    Rasterise each page at the target DPI/JPEG quality.
    screen=72dpi/20%, ebook=96dpi/45%, printer=150dpi/80%, prepress=200dpi/90%.
    """
    PRESET = {
        'screen':   (72,  20),
        'ebook':    (96,  45),
        'printer':  (150, 80),
        'prepress': (200, 90),
    }
    dpi, jpeg_q = PRESET.get(quality, (96, 45))
    scale = dpi / 72
    mat = fitz.Matrix(scale, scale)

    src = fitz.open(input_path)
    out = fitz.open()
    for page in src:
        pix = page.get_pixmap(matrix=mat, alpha=False)
        buf = io.BytesIO()
        img = Image.frombytes('RGB', [pix.width, pix.height], pix.samples)
        img.save(buf, 'JPEG', quality=jpeg_q, optimize=True)
        buf.seek(0)
        new_page = out.new_page(width=page.rect.width, height=page.rect.height)
        new_page.insert_image(new_page.rect, stream=buf.read())
    out.save(output_path, garbage=4, deflate=True)
    src.close()
    out.close()
    return output_path

# ── rotate ────────────────────────────────────────────────────────────────────

def rotate_pdf(input_path: Path, output_path: Path, degrees: int, pages: list[int] | None = None) -> Path:
    with fitz.open(input_path) as doc:
        targets = pages if pages else list(range(1, len(doc) + 1))
        for p_num in targets:
            if 1 <= p_num <= len(doc):
                page = doc[p_num - 1]
                page.set_rotation((page.rotation + degrees) % 360)
        doc.save(output_path)
    return output_path

# ── pdf → docx (pdfminer.six + python-docx) ──────────────────────────────────

from .image_operations import convert_image, crop_image, resize_image

from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer, LTChar
from docx import Document
from docx.shared import Pt


def pdf_to_docx(input_path: Path, output_path: Path) -> Path:
    doc = Document()

    for page_layout in extract_pages(str(input_path)):
        blocks = []
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                text = element.get_text().strip()
                if text:
                    font_size = 11
                    for text_line in element:
                        for char in text_line:
                            if isinstance(char, LTChar):
                                font_size = round(char.size)
                                break
                        break
                    blocks.append((element.y1, text, font_size))

        blocks.sort(key=lambda b: -b[0])

        for _, text, font_size in blocks:
            if font_size >= 14:
                doc.add_heading(text, level=1)
            elif font_size >= 12:
                doc.add_heading(text, level=2)
            else:
                para = doc.add_paragraph(text)
                para.style.font.size = Pt(font_size)

        doc.add_page_break()

    doc.save(str(output_path))
    return output_path


def pdf_to_docx_ocr(input_path: Path, output_path: Path, lang: str = 'por+eng') -> Path:
    import ocrmypdf
    languages = lang.replace('+', ',').split(',')
    with tempfile.TemporaryDirectory() as tmpdir:
        ocr_pdf_path = Path(tmpdir) / 'ocr_input.pdf'
        ocrmypdf.ocr(
            str(input_path),
            str(ocr_pdf_path),
            language=languages,
            skip_text=True,
            deskew=True,
            progress_bar=False,
        )
        pdf_to_docx(ocr_pdf_path, output_path)
    return output_path


# ── docx → pdf (LibreOffice) ──────────────────────────────────────────────────

def _soffice_convert(input_path: Path, output_path: Path) -> None:
    """
    Convert a DOCX/DOC to PDF via LibreOffice headless.
    Profile dir is kept outside the conversion tmpdir to avoid I/O lock
    conflicts (0xc10) when soffice writes its lock files alongside the output.
    """
    soffice = resolve_executable('soffice')

    profile_dir = Path(tempfile.gettempdir()) / f'lo_{uuid.uuid4().hex}'
    profile_dir.mkdir(parents=True, exist_ok=True)

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            tmp_input = tmpdir_path / f'input{input_path.suffix}'
            shutil.copy2(input_path, tmp_input)

            env = os.environ.copy()
            env['HOME'] = str(Path.home())
            if 'DISPLAY' not in env:
                env['DISPLAY'] = ':0'

            cmd = [
                str(soffice),
                f'-env:UserInstallation=file://{profile_dir}',
                '--headless',
                '--norestore',
                '--convert-to', 'pdf',
                '--outdir', str(tmpdir_path),
                str(tmp_input),
            ]

            result = subprocess.run(cmd, capture_output=True, text=True, timeout=600, env=env)

            if result.returncode != 0:
                raise RuntimeError(
                    f'LibreOffice falhou (rc={result.returncode}): {result.stderr.strip()}'
                )

            generated = tmpdir_path / f'input.pdf'
            if not generated.exists():
                raise FileNotFoundError(
                    f'LibreOffice não gerou o arquivo esperado. '
                    f'Tmpdir: {list(tmpdir_path.iterdir())}'
                )

            shutil.move(str(generated), str(output_path))
    finally:
        shutil.rmtree(profile_dir, ignore_errors=True)


def docx_to_pdf(input_path: Path, output_path: Path) -> Path:
    _soffice_convert(input_path, output_path)
    return output_path


# ── dispatcher ────────────────────────────────────────────────────────────────

def execute_job(job, workspace: Path) -> str:
    """
    Resolve input paths, run the operation, return output filename.
    Used by both the sync view path and the Huey task.
    """
    from common.executables import safe_path

    in_paths = [safe_path(workspace, f) for f in job.input_files]
    params = job.params or {}
    op = job.operation
    jid = job.id

    if op == 'merge':
        out = workspace / f'merged_{jid}.pdf'
        merge_pdfs(in_paths, out)
        return out.name

    if op == 'split':
        raw_ranges = params.get('ranges', [])
        ranges = [tuple(r) for r in raw_ranges]
        tmp_dir = workspace / f'split_{jid}_parts'
        parts = split_pdf(in_paths[0], tmp_dir, ranges)
        out = workspace / f'split_{jid}.zip'
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zf:
            for p in parts:
                zf.write(p, p.name)
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return out.name

    if op == 'pdf_to_images':
        dpi = int(params.get('dpi', 150))
        tmp_dir = workspace / f'images_{jid}'
        imgs = pdf_to_images(in_paths[0], tmp_dir, dpi)
        out = workspace / f'images_{jid}.zip'
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zf:
            for p in imgs:
                zf.write(p, p.name)
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return out.name

    if op == 'images_to_pdf':
        out = workspace / f'combined_{jid}.pdf'
        images_to_pdf(in_paths, out)
        return out.name

    if op == 'pdf_to_text':
        out = workspace / f'text_{jid}.txt'
        pdf_to_text(in_paths[0], out)
        return out.name

    if op == 'ocr_image':
        lang = params.get('lang', 'por+eng')
        out = workspace / f'ocr_{jid}.txt'
        ocr_image(in_paths[0], out, lang)
        return out.name

    if op == 'ocr_pdf':
        lang = params.get('lang', 'por+eng')
        out = workspace / f'ocr_{jid}.pdf'
        ocr_pdf(in_paths[0], out, lang)
        return out.name

    if op == 'compress':
        quality = params.get('quality', 'ebook')
        out = workspace / f'compressed_{jid}.pdf'
        compress_pdf(in_paths[0], out, quality)
        return out.name

    if op == 'rotate':
        degrees = int(params.get('degrees', 90))
        pages = params.get('pages') or None
        out = workspace / f'rotated_{jid}.pdf'
        rotate_pdf(in_paths[0], out, degrees, pages)
        return out.name

    if op == 'pdf_to_docx':
        out = workspace / f'converted_{jid}.docx'
        pdf_to_docx(in_paths[0], out)
        return out.name

    if op == 'pdf_to_docx_ocr':
        lang = params.get('lang', 'por+eng')
        out = workspace / f'converted_{jid}.docx'
        pdf_to_docx_ocr(in_paths[0], out, lang)
        return out.name

    if op == 'docx_to_pdf':
        out = workspace / f'converted_{jid}.pdf'
        docx_to_pdf(in_paths[0], out)
        return out.name

    if op == 'img_convert':
        fmt    = params.get('output_format', 'png')
        quality = int(params.get('quality', 90))
        out = workspace / (Path(in_paths[0]).stem + f'.{fmt}')
        convert_image(in_paths[0], out, fmt, quality)
        return out.name

    if op == 'img_crop':
        x = int(params.get('x', 0))
        y = int(params.get('y', 0))
        w = int(params.get('width', 100))
        h = int(params.get('height', 100))
        out = workspace / (Path(in_paths[0]).stem + '_crop' + Path(in_paths[0]).suffix)
        crop_image(in_paths[0], out, x, y, w, h)
        return out.name

    if op == 'img_resize':
        width  = params.get('width')
        height = params.get('height')
        keep   = bool(params.get('keep_aspect', True))
        out = workspace / (Path(in_paths[0]).stem + '_resized' + Path(in_paths[0]).suffix)
        resize_image(in_paths[0], out, int(width) if width else None, int(height) if height else None, keep)
        return out.name

    raise ValueError(f'Operação desconhecida: {op}')
